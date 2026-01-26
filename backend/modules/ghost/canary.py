"""
Canary Token Generator
Creates bait files that alert when opened.
"""

import os
import uuid
import json
from datetime import datetime
from typing import Dict, Optional

GHOST_DIR = os.path.join(os.path.dirname(__file__), "../../../ghost_data")
ALERTS_DIR = os.path.join(GHOST_DIR, "alerts")
os.makedirs(ALERTS_DIR, exist_ok=True)


def generate_canary_token(token_id: str, file_type: str = "pdf", metadata: Dict = None) -> str:
    """
    Create a canary token file of the specified type and persist its metadata.
    
    Parameters:
        token_id (str): Identifier for the token; used as the token directory name and file basename.
        file_type (str): Desired file type to generate. Supported values: "pdf", "docx", "html". Any other value creates a plain text file.
        metadata (Dict, optional): Arbitrary metadata to store alongside the token; stored in token.json.
    
    Returns:
        file_path (str): Filesystem path to the created token file.
    
    Raises:
        RuntimeError: If `file_type` is "docx" and the required `python-docx` package is not installed.
    """
    token_dir = os.path.join(GHOST_DIR, "tokens", token_id)
    os.makedirs(token_dir, exist_ok=True)
    
    # Get callback URL (would be configured in production)
    callback_url = os.getenv("CANARY_CALLBACK_URL", f"http://127.0.0.1:8000/api/v1/trapdoor/ghost/trap/{token_id}")
    
    # Create token file based on type
    if file_type == "pdf":
        file_path = os.path.join(token_dir, f"{token_id}.pdf")
        # Create minimal PDF with tracking pixel/URL
        # In production, would embed invisible tracking
        with open(file_path, "wb") as f:
            f.write(b"%PDF-1.4\n")  # Minimal PDF header
    elif file_type == "docx":
        file_path = os.path.join(token_dir, f"{token_id}.docx")
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to generate DOCX canary files.") from exc

        doc = Document()
        doc.add_heading("Document", level=1)
        doc.add_paragraph("This document contains sensitive information.")
        doc.add_paragraph(f"Tracking URL: {callback_url}")
        doc.save(file_path)
    elif file_type == "html":
        file_path = os.path.join(token_dir, f"{token_id}.html")
        # Create HTML file with hidden image beacon
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Document</title>
</head>
<body>
    <h1>Document</h1>
    <p>This document contains sensitive information.</p>
    <!-- Hidden tracking pixel -->
    <img src="{callback_url}" width="1" height="1" style="display:none;" />
</body>
</html>"""
        with open(file_path, "w") as f:
            f.write(html_content)
    else:
        file_path = os.path.join(token_dir, f"{token_id}.txt")
        with open(file_path, "w") as f:
            f.write(f"Canary Token: {token_id}\n")
    
    # Store token metadata
    token_meta = {
        "token_id": token_id,
        "file_type": file_type,
        "created_at": datetime.now().isoformat(),
        "metadata": metadata or {},
        "file_path": file_path
    }
    
    meta_path = os.path.join(token_dir, "token.json")
    with open(meta_path, "w") as f:
        json.dump(token_meta, f, indent=2)
    
    return file_path


def check_canary_alert(token_id: str) -> Optional[Dict]:
    """Check if canary token was triggered."""
    alert_path = os.path.join(ALERTS_DIR, f"{token_id}.json")
    
    if os.path.exists(alert_path):
        with open(alert_path, "r") as f:
            return json.load(f)
    
    return None


def record_canary_alert(token_id: str, ip_address: str, user_agent: str, referrer: str = None):
    """Record a canary token alert."""
    alert = {
        "token_id": token_id,
        "triggered_at": datetime.now().isoformat(),
        "ip_address": ip_address,
        "user_agent": user_agent,
        "referrer": referrer
    }
    
    alert_path = os.path.join(ALERTS_DIR, f"{token_id}.json")
    with open(alert_path, "w") as f:
        json.dump(alert, f, indent=2)
    
    return alert